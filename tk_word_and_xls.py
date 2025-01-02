from email.policy import default
from tkinter import *
from tkinter import filedialog
from docx import Document
import openpyxl
from docx.enum.text import WD_ALIGN_PARAGRAPH


def open_file():
    filepath=filedialog.askopenfilename()
    if filepath != "":
        #doc.docx=filepath
        doc = Document(docx=filepath)
        path.value=filepath
        print(path.value)
        table = doc.tables[1]
        #print(table.cell(1, 2).text)
        val_old.value=table.cell(1, 2).text
        print(val_old.value)

def open_table():
    tablepath=filedialog.askopenfilename()
    if tablepath !="":
        wb = openpyxl.load_workbook(tablepath)
        sheet=wb.active
        val_new.value=str(sheet['C25'].value)
        print(val_new.value)

def save_new_file():
    val_old.set(val_new.get())
    #val_old.value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    print(path.value)
    doc=Document(docx=path.value)
    table = doc.tables[1]
    print(table.cell(1, 2).text)
    table.cell(1, 2).text = val_new.value
    table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(path.value)
    print(table.cell(1, 2).text)

root=Tk()
root.title('Приложение Систерм')
root.geometry('400x400+200+200')
root.iconbitmap(default='brend.ico')

val_old=StringVar()
val_new=StringVar()
path=StringVar()


file_button=Button(text='Открыть отчет', command=open_file)
file_button.place(x=20, y=20)
btn=Button(text='Открыть таблицу', command=open_table)
btn.place(x=20, y=60)
btn_safe=Button(text='Изменить', command=save_new_file)
btn_safe.place(x=20, y=120)

root.mainloop()



