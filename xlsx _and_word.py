
from docx import Document
import openpyxl
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(docx = 'Инструкция-по-автоматике-Пороховой-ГВС-вода.docx')
table=doc.tables[1]
print(table.cell(1, 2).text)
wb = openpyxl.load_workbook('Сетевые_Порох3_вода.xlsx')
sheet = wb.active
print(sheet['A1'].value)
table.cell(1, 2).text= str(sheet['C25'].value)
table.cell(1, 2).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT
doc.save('Инструкция-по-автоматике-Пороховой-ГВС-вода.docx')
print(table.cell(1, 2).text)