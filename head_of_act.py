from docx import Document
from tkinter import *
from tkinter import filedialog

def open_file():
    filepath=filedialog.askopenfilename()
    global document
    if filepath != "":
        doc = Document(docx=filepath)

        text=[]
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        print('\n'.join(text))

        new_docs=doc.paragraphs.copy()
        for new_doc in new_docs:
            document.add_paragraph(new_doc.text)


def save_new_file():
    filepath = filedialog.asksaveasfilename()
    if filepath != "":
        document.save(filepath)





document=Document()

root=Tk()
root.title('Приложение Систерм')
root.geometry('400x400+200+200')
root.iconbitmap(default='brend.ico')

file_button=Button(text='Открыть файл', command=open_file)
file_button.place(x=20, y=20)

btn_safe=Button(text='Сохранить файл', command=save_new_file)
btn_safe.place(x=20, y=120)

root.mainloop()
