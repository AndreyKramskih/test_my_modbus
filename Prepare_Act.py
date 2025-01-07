import numpy as np
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from tkinter import *
from tkinter import filedialog
from docx import Document
import openpyxl
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook

def safe_act():
    filepath = filedialog.asksaveasfilename()
    if filepath != "":
        document.save(filepath)

def open_table():
    table_path=filedialog.askopenfilename()
    if table_path !="":

        #wb = openpyxl.load_workbook(table_path)
        #sheet=wb.active
        #val=str(sheet['A1'].value)
        #print(val)
        #df = pd.read_excel(table_path, sheet_name='Table 1', index_col=0, skiprows=2)

        df = pd.read_excel(table_path, sheet_name='Table 1', skiprows=2)
        df_cleaned=df.dropna()
        print(df_cleaned.head())
        print(len(df_cleaned.index))
        print(df_cleaned.columns)
        print(len(df_cleaned.to_numpy()))
        for i in range(0,5):
            print(df_cleaned.to_numpy()[i])
        xl_arr=df_cleaned.to_numpy()

        xl_list=xl_arr.tolist()
        print(xl_list[0][2])
        print(xl_list[2])
        j=1


        for i in range(0, len(xl_arr)):
                xl_list[i].insert(0,j)
                j+=1
        j=1



        f_list=np.empty((1,len(xl_arr.tolist()[0]))).tolist()
        f_list.clear()
        #print(f_list)
        s_list = np.empty((1, len(xl_arr.tolist()[0]))).tolist()
        s_list.clear()
        th_list = np.empty((1, len(xl_arr.tolist()[0]))).tolist()
        th_list.clear()

        f_list += [x for x in xl_arr.tolist() if 'теплообменник' in str(x).lower()]
        f_list += [x for x in xl_arr.tolist() if 'насос' in str(x).lower()]
        f_list += [x for x in xl_arr.tolist() if 'регулирующий' in str(x).lower()]
        f_list += [x for x in xl_arr.tolist() if 'регулятор давления' in str(x).lower()]

        for i in range(0, len(f_list)):
                f_list[i].insert(0,j)
                j+=1
        j=1
        #print(f_list)
        s_list += [x for x in xl_arr.tolist() if 'фильтр' in str(x).lower()]
        s_list += [x for x in xl_arr.tolist() if 'обратный' in str(x).lower()]
        s_list += [x for x in xl_arr.tolist() if 'вентиль' in str(x).lower()]
        s_list += [x for x in xl_arr.tolist() if 'шаровой' in str(x).lower()]

        for i in range(0, len(s_list)):
            s_list[i].insert(0, j)
            j += 1
        j = 1

        th_list += [x for x in xl_arr.tolist() if 'Манометр' in str(x)]
        th_list += [x for x in xl_arr.tolist() if 'термометр' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'термостат погружной' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'датчик' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'реле' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'прессостат' in str(x).lower()]

        for i in range(0, len(th_list)):
            th_list[i].insert(0, j)
            j += 1
        j = 1

        headers = ('№ ', 'Поз.', 'Наименование', 'Тип, марка\nматериал\nТехническая\nдокументация',
                   'Завод -\nизготовитель', 'Кол-\nво,\nшт')

        global document
        #document.add_heading('Таблица 1')
        #table1 = create_table(document,headers, xl_list)
        #document.add_paragraph()
        document.add_heading('Таблица 1')
        table2 = create_table(document, headers, f_list)
        document.add_paragraph()
        document.add_heading('Таблица 2')
        table3 = create_table(document, headers, s_list)
        document.add_paragraph()
        document.add_heading('Таблица 3')
        table4 = create_table(document, headers, th_list)


def create_table(document, headers, rows, style='Table Grid'):
    cols_number = len(headers)

    table = document.add_table(rows=1, cols=cols_number)
    table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])

    return table


document = Document()

#headers = ('№ ', 'Поз.', 'Наименование', 'Тип, марка\nматериал', 'Техническая\nдокументация',
          # 'Завод -\nизготовитель', 'Кол-\nво,\nшт')
#records_table1 = (
    #(0, 'Nan', 'Nan', 0, 2, 3, 4),
    #(0, 'Nan', 'Nan', 0, 2, 3, 4),
    #(0, 'Nan', 'Nan', 0, 2, 3, 4),
    #(0, 'Nan', 'Nan', 0, 2, 3, 4)
#)
#table1 = create_table(document, headers, records_table1)

#document.add_paragraph()

#rows = [
    #[x, x, x * x] for x in range(1, 10)
#]
#table2 = create_table(document, ('x', 'y', 'x * y'), rows)

#document.save('C:/Users/Andrey Kramskikh/Downloads/Акт_тест.docx')


root=Tk()
root.title('Приложение Систерм')
root.geometry('400x400+200+200')
root.iconbitmap(default='brend.ico')

file_button=Button(text='Открыть спец', command=open_table)
file_button.place(x=20, y=20)
btn=Button(text='Создать отчет', command=safe_act)
btn.place(x=20, y=60)


root.mainloop()