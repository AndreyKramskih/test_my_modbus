#import aspose.words as aw
from xml.dom.minidom import Document

import docx
# load Word document
#doc = aw.Document("Инструкция  по автоматике Пороховой ГВС вода.doc")

# replace text
#doc.range.replace("ГВС", "Вентиляция", aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))

# save the modified document
#doc.save("Инструкция  по автоматике Пороховой Вентиляция вода.doc")

from docx import Document
doc = Document(docx = 'Инструкция-по-автоматике-Пороховой-ГВС-вода.docx')

paras = doc.paragraphs

name = "Вентиляция"

#for para in paras:
   # para.text = para.text.replace('ГВС', name)


for para in paras:
    for run in para.runs:
        if 'ГВС' in run.text:
            run.text = run.text.replace('ГВС', name)
doc.save('Инструкция  по автоматике Пороховой Вентиляция вода.docx')