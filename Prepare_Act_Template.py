import numpy as np
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from tkinter import *
from tkinter import filedialog
from docx import Document
import openpyxl
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from docxtpl import DocxTemplate

def safe_act():
    filepath = filedialog.asksaveasfilename()
    if filepath != "":
        document.save(filepath)

def open_table():
    table_path=filedialog.askopenfilename()
    if table_path !="":
        df = pd.read_excel(table_path, sheet_name='Table 1', skiprows=2)
        df_cleaned=df.dropna()
        print(df_cleaned.head())
        print(len(df_cleaned.index))
        print(df_cleaned.columns)
        print(len(df_cleaned.to_numpy()))
        for i in range(0,5):
            print(df_cleaned.to_numpy()[i])
        xl_arr=df_cleaned.to_numpy()

        xl_list=xl_arr.tolist()
        print(xl_list[0][2])
        print(xl_list[2])
        j=1

        for i in range(0, len(xl_arr)):
                xl_list[i].insert(0,j)
                j+=1
        j=1

        f_list=np.empty((1,len(xl_arr.tolist()[0]))).tolist()
        f_list.clear()
        #print(f_list)
        #s_list = np.empty((1, len(xl_arr.tolist()[0]))).tolist()
        #s_list.clear()
        #th_list = np.empty((1, len(xl_arr.tolist()[0]))).tolist()
        #th_list.clear()

        f_list += [x for x in xl_arr.tolist() if 'теплообменник' in str(x).lower()]
        f_list += [x for x in xl_arr.tolist() if 'насос' in str(x).lower()]
        f_list += [x for x in xl_arr.tolist() if 'регулирующий' in str(x).lower()]
        f_list += [x for x in xl_arr.tolist() if 'регулятор давления' in str(x).lower()]

        for i in range(0, len(f_list)):
                f_list[i].insert(0,j)
                j+=1
        j=1

        """""
        s_list += [x for x in xl_arr.tolist() if 'фильтр' in str(x).lower()]
        s_list += [x for x in xl_arr.tolist() if 'обратный' in str(x).lower()]
        s_list += [x for x in xl_arr.tolist() if 'вентиль' in str(x).lower()]
        s_list += [x for x in xl_arr.tolist() if 'шаровой' in str(x).lower()]

        for i in range(0, len(s_list)):
            s_list[i].insert(0, j)
            j += 1
        j = 1

        th_list += [x for x in xl_arr.tolist() if 'Манометр' in str(x)]
        th_list += [x for x in xl_arr.tolist() if 'термометр' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'термостат погружной' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'датчик' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'реле' in str(x).lower()]
        th_list += [x for x in xl_arr.tolist() if 'прессостат' in str(x).lower()]

        for i in range(0, len(th_list)):
            th_list[i].insert(0, j)
            j += 1
        j = 1
        """""
        headers = ('№ ', 'Поз.', 'Наименование', 'Тип, марка\nматериал\nТехническая\nдокументация',
                   'Завод -\nизготовитель', 'Кол-\nво,\nшт')

        global document

        all_tables = document.tables
        new_table = all_tables[0]
        print(new_table)
        cols_number = len(headers)
        for row in f_list:
            row_cells = new_table.add_row().cells
            for i in range(cols_number):
               row_cells[i].text = str(row[i])

        #document.add_heading('Таблица 1')
        #table1 = create_table(document,headers, xl_list)
        #document.add_paragraph()
        #document.add_heading('Таблица 1')
        #table2 = create_table(document, headers, f_list)
        document.add_paragraph()
        #document.add_heading('Таблица 2')
        #table3 = create_table(document, headers, s_list)
        #document.add_paragraph()
        #document.add_heading('Таблица 3')
        #table4 = create_table(document, headers, th_list)


def create_table(document, headers, rows, style='Table Grid'):
    cols_number = len(headers)

    table = document.add_table(rows=1, cols=cols_number)
    table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])

    return table


#document = Document()

# Загрузка шаблона
document = DocxTemplate("Шаблон.docx")

# Данные для заполнения шаблона
context = {
    'station': 'Блочная установка узел смешения гликоля ',
    'calc': 'GEVHeat TPZ-ИТП-600.294',
    'company': 'Энергия Технологий',
    'object': 'Реконструкция поверхностного комплекса НШ-1 НШПП «Яреганефть» АО «Соликамский завод УРАЛ».',
    'address': 'РФ, Ярегское месторождение',
    'number': '1.1',
    'name': 'ВХОДНОГО КОНТРОЛЯ ОБОРУДОВАНИЯ',
    'data': '02.07.2024'
}

# Заполнение шаблона данными
document.render(context)

# Сохранение документа
#doc.save("новый_документ.docx")

root=Tk()
root.title('Подготовка Актов')
root.geometry('600x600+200+200')
root.iconbitmap(default='brend.ico')

file_button=Button(text='Открыть спец', command=open_table)
file_button.place(x=20, y=20)
btn=Button(text='Создать отчет', command=safe_act)
btn.place(x=20, y=60)

root.mainloop()